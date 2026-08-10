#!/usr/bin/env python3
"""Build student and teacher DOCX versions of the e网通 timed training."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "20出题" / "202607暑假高三-e网通课堂限时训练.md"
STUDENT = ROOT / "20出题" / "202607暑假高三-e网通课堂限时训练-学生版.docx"
TEACHER = ROOT / "20出题" / "202607暑假高三-e网通课堂限时训练-答案解析.docx"

# compact_reference_guide with named Chinese exam overrides:
# A4 portrait; 15 mm margins; body 宋体 10.5 pt, 1.25 lines; restrained blue headings.
ACCENT = RGBColor(31, 78, 121)
MUTED = RGBColor(89, 89, 89)
LIGHT = "E8EEF5"


def set_run_font(run, name: str, size: float, *, bold: bool = False, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri" if name != "Consolas" else "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri" if name != "Consolas" else "Consolas")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def shade_paragraph(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run_font(run, "宋体", 9, color=MUTED)


def configure_styles(doc: Document):
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(15)
    sec.bottom_margin = Mm(15)
    sec.left_margin = Mm(16)
    sec.right_margin = Mm(16)
    sec.header_distance = Mm(7)
    sec.footer_distance = Mm(7)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, before, after in (
        ("Heading 1", 14, 10, 6),
        ("Heading 2", 12, 8, 4),
        ("Heading 3", 11.5, 6, 3),
    ):
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.add_run("202607 暑假高三 · e网通精选"), "宋体", 8.5, color=MUTED)
    add_page_field(sec.footer.paragraphs[0])


def add_inline(paragraph, text: str, *, bold=False, size=10.5):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, "Consolas", 9.5, bold=bold)
        else:
            clean = re.sub(r"\[\[([^]|]+/)?([^]|]+)(?:\|[^]]+)?\]\]", r"\2", part)
            run = paragraph.add_run(clean)
            set_run_font(run, "宋体", size, bold=bold)


def add_image(doc: Document, relative: str):
    path = ROOT / relative.replace("/", "\\")
    from PIL import Image

    with Image.open(path) as im:
        width_px, height_px = im.size
    max_w = Cm(15.5)
    max_h = Cm(8.0)
    ratio = width_px / height_px
    width = max_w
    height = width / ratio
    if height > max_h:
        height = max_h
        width = height * ratio
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_together = True
    p.add_run().add_picture(str(path), width=width, height=height)


def add_code(doc: Document, lines: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.right_indent = Cm(0.2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.0
    # Short snippets should stay together; the long Q5 program must be allowed
    # to flow across pages instead of being compressed or clipped.
    p.paragraph_format.keep_together = len(lines) <= 24
    shade_paragraph(p, "F2F4F7")
    run = p.add_run("\n".join(lines))
    set_run_font(run, "Consolas", 8.5)


def add_metadata_strip(doc: Document):
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [Cm(5.9), Cm(5.9), Cm(5.9)]
    values = [("45 分钟", "限时"), ("31 分", "满分"), ("5 题", "题量")]
    for cell, width, (value, label) in zip(table.rows[0].cells, widths, values):
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=100, bottom=100, start=100, end=100)
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), LIGHT)
        tc_pr.append(shd)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(value), "黑体", 11, bold=True, color=ACCENT)
        set_run_font(p.add_run(f"\n{label}"), "宋体", 8.5, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def render_markdown(doc: Document, markdown: str, include_answers: bool):
    markdown = markdown.lstrip("\ufeff")
    if markdown.startswith("---\n"):
        end = markdown.find("\n---\n", 4)
        if end != -1:
            markdown = markdown[end + 5 :].lstrip()
    if not include_answers:
        markdown = re.split(
            r"\n---\n\s*\n## (?:答案|参考答案与解析)\s*$",
            markdown,
            maxsplit=1,
            flags=re.MULTILINE,
        )[0]

    lines = markdown.splitlines()
    in_code = False
    code_lines: list[str] = []
    title_done = False
    skip_metadata_line = False

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                add_code(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line:
            continue
        if line == "---":
            doc.add_page_break()
            continue
        image_match = re.fullmatch(r"!\[\[(attachments/[^]]+)\]\]", line)
        if image_match:
            add_image(doc, image_match.group(1))
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            set_run_font(p.add_run(line[2:]), "黑体", 18, bold=True, color=ACCENT)
            title_done = True
            continue
        if title_done and line.startswith("时间："):
            add_metadata_strip(doc)
            continue
        if title_done and line.startswith("建议用时："):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            add_inline(p, line, size=8.5)
            for run in p.runs:
                run.font.color.rgb = MUTED
            continue
        if title_done and line.startswith("姓名："):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(6)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, line, size=10.5)
            title_done = False
            continue
        if line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, line[3:], bold=True, size=14)
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[4:], bold=True, size=12)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.first_line_indent = Cm(-0.15)
            p.paragraph_format.space_after = Pt(1)
            add_inline(p, line[2:])
            continue
        p = doc.add_paragraph()
        p.paragraph_format.widow_control = True
        if line.startswith("出处："):
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(6)
            add_inline(p, line, size=9)
            for run in p.runs:
                run.font.color.rgb = MUTED
        else:
            add_inline(p, line)


def build(target: Path, include_answers: bool):
    doc = Document()
    configure_styles(doc)
    markdown = SOURCE.read_text(encoding="utf-8")
    render_markdown(doc, markdown, include_answers)
    doc.core_properties.title = "202607 暑假高三信息技术课堂限时训练"
    doc.core_properties.subject = "e网通六卷精选，45分钟"
    doc.core_properties.author = "慈溪中学信息技术组"
    try:
        doc.save(target)
        actual_target = target
    except PermissionError:
        actual_target = target.with_name(f"{target.stem}-更新{target.suffix}")
        doc.save(actual_target)
    print(actual_target)
    return actual_target


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--student-only",
        action="store_true",
        help="Build only the student version.",
    )
    args = parser.parse_args()
    build(STUDENT, include_answers=False)
    if not args.student_only:
        build(TEACHER, include_answers=True)


if __name__ == "__main__":
    main()
