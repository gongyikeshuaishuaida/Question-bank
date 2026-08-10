from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT_DIR = Path("20出题")
STUDENT_DOCX = OUT_DIR / "202606九校期末错题专题练习卷.docx"
ANSWER_DOCX = OUT_DIR / "202606九校期末错题专题练习卷-答案解析.docx"


def set_font(run, size=10.5, bold=False, name="宋体"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def new_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(1.7)
    sec.right_margin = Cm(1.7)
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)
    return doc


def para(doc, text="", first=False):
    text = text.replace("`", "")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if first:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_font(r)
    return p


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_font(r, 16, True, "黑体")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_font(r, 9)


def section(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(text)
    set_font(r, 12, True, "黑体")


def code(doc, text):
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.62)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        set_font(r, 9.2, False, "Consolas")


def build_student():
    doc = new_doc()
    add_title(
        doc,
        "202606 九校期末错题专题提升卷（难题精选）",
        "题库精选：pandas 筛选/分组/排序 + 滑动平均/滑动窗口/单调队列",
    )
    para(doc, "班级：__________  姓名：__________  建议用时：55 分钟")

    section(doc, "一、pandas 综合处理")
    para(doc, "1. 某用电监测系统采集了部分房间 2025 年的照明用电数据。现要找出月均照明用电量最多的房间，并绘制该房间各月份照明用电量折线图，再按用电量降序输出该房间各月份用电情况。", True)
    code(doc, '''
import pandas as pd
import matplotlib.pyplot as plt
df = pd.read_excel("data.xlsx")
columns = df.columns[1:]
maxv = 0
for col in columns:
    s = df[col].sum()
    ave = round(s / 12, 2)
    if ave > maxv:
        ____①____
        maxcol = col
df.insert(0, "月", "")
for i in df.index:
    t = df.at[i, "时间"]
    df.at[i, "月"] = ____②____
df = df.drop("时间", axis=1)
df2 = df[["月", maxcol]]
____③____
____④____
plt.show()
____⑤____
for i in df3.index:
    print(df3.at[i, "月"], round(df3[maxcol][i], 2))
''')
    para(doc, "③④⑤可选：A. `df3=df.groupby(\"月\", as_index=True).mean()`  B. `df3=df2.groupby(\"月\", as_index=False).sum()`")
    para(doc, "C. `df3=df3.sort_values(maxcol, ascending=False)`  D. `df3=df3.sort_values(col, ascending=False)`")
    para(doc, "E. `plt.plot(df3.index, df3[maxcol])`  F. `plt.plot(df3.月, df3[maxcol])`")
    para(doc, "答：①____________________  ②____________________  ③______  ④______  ⑤______")

    para(doc, "2. 统计上一年总客流量最大的通道，并根据该通道月平均客流量最大的五个月份绘制柱形图。请选择合适代码填入①②③处。", True)
    code(doc, '''
import pandas as pd
import matplotlib.pyplot as plt
df = pd.read_excel("data.xlsx")
df1 = _____①_____
df1 = df1.sort_values("通道客流量", ascending=False)
c = df1.at[0, "通道名称"]
df2 = _____②_____
df3 = df2.groupby("月份", as_index=False).mean()
df3 = _____③_____
df3 = df3.tail(5)
plt.bar(df3.月份, df3.通道客流量)
plt.show()
''')
    para(doc, "A. `df.groupby(\"月份\", as_index=False).sum()`  B. `df.groupby(\"通道名称\", as_index=False).sum()`")
    para(doc, "C. `df[df.通道名称 == c]`  D. `df1[df1.通道名称 == c]`")
    para(doc, "E. `df3.sort_values(\"通道客流量\", ascending=False)`  F. `df3.sort_values(\"通道客流量\", ascending=True)`")
    para(doc, "答：①______  ②______  ③______")

    para(doc, "3. 宿舍用电功率超过 2000W 时发送警报。统计本周引发警报次数最多的宿舍，并绘制该宿舍每天警报次数折线图。请选择合适代码填入①②③处。", True)
    code(doc, '''
df = pd.read_excel("data.xlsx")
df1 = ____①____
df2 = df1.groupby("宿舍号", as_index=True).count()
# 重命名 df2 中“功率”列名称为“次数”，代码略
df3 = ____②____
m = df3.index[0]
df4 = df1[df1["宿舍号"] == m]
# 根据 df4["时间"] 提取“日”列，代码略
df5 = ____③____
plt.plot(df5.index, df5.values)
''')
    para(doc, "A. `df['功率'] > 2000`  B. `df[df['功率'] > 2000]`")
    para(doc, "C. `df2.sort_values('功率', ascending=False)`  D. `df2.sort_values('次数', ascending=False)`")
    para(doc, "E. `df4.groupby('日', as_index=False).count()`  F. `df4.groupby('日', as_index=True)['宿舍号'].count()`")
    para(doc, "答：①______  ②______  ③______")

    para(doc, "4. 统计每天噪声大于 55 分贝的次数，并输出次数最多的前 3 天。请选择合适代码填入①②③处。", True)
    code(doc, '''
df = pd.read_excel("noise.xlsx")
df1 = ____①____
df2 = ____②____
df2 = df2.rename(columns={"噪声": "次数"})
df3 = ____③____
df4 = df3.head(3)
print(df4)
''')
    para(doc, "A. `df[df.噪声>55]`  B. `df[df[噪声]>55]`  C. `df1.groupby(\"日期\", as_index=False).count()`")
    para(doc, "D. `df1.groupby(\"日期\", as_index=False).sum()`  E. `df2.sort_values(\"次数\")`  F. `df2.sort_values(\"次数\", ascending=False)`")
    para(doc, "答：①______  ②______  ③______")

    section(doc, "二、滑动平均、窗口边界、单调队列")
    para(doc, "5. 实时温度以任意连续 k 个数据为一组计算平均值；若连续出现 m 个平均值均大于阈值 pt，则判定温度过高。", True)
    para(doc, "（1）温度序列为 49.90, 50.10, 49.80, 50.20, 49.90, 50.20，k=3，pt=50，则大于阈值 pt 的平均值有 ______ 个。")
    code(doc, '''
lst = [0.0] * k
i = 0
c = 0
____①____
while True:
    # 获取实时温度 tmp，代码略
    v = i % k
    total = total + tmp - lst[v]
    ____②____
    if i >= k - 1:
        ave = total / k
        if ____③____:
            c = c + 1
        else:
            c = 0
        if c >= m:
            # 发出警报，代码略
            c = 0
    i = i + 1
''')
    para(doc, "答：（1）______；①____________________  ②____________________  ③____________________")

    para(doc, "6. 若连续 4 次湿度平均值在 70±5 范围内，且波动不超过 5，则为稳定状态，否则为波动状态。", True)
    para(doc, "（1）初始为稳定状态，最近 7 次湿度为 [64,66,66,68,72,68,68]，期间修改 ______ 次检测频率。")
    code(doc, '''
hum = [64, 66, 66, 68, 72, 68, 68]
m = 4
s = 70
T = 5
cnt = 0
pre_state = "1 次/分钟(稳定状态)"
for i in range(0, ____①____):
    cur = hum[i : i + m]
    tot = 0
    for num in cur:
        tot += num
    avg = tot / len(cur)
    bd = max(cur) - min(cur)
    if s - T <= avg <= s + T and ____②____:
        state = "1 次/分钟(稳定状态)"
    else:
        state = "5 次/分钟(波动状态)"
    if ____③____:
        cnt += 1
    pre_state = state
''')
    para(doc, "答：（1）______；①____________________  ②____________________  ③____________________")

    para(doc, "7. 每 5 分钟采集一次，flag[i] 表示第 i 个时刻是否异常。找出异常点最多的连续 1 小时时段，数量相同取最早。", True)
    code(doc, '''
n = len(a)
flag = [0] * n
for i in range(n):
    if a[i][2] > std:
        flag[i] = 1
start_time = cnt = 0
for i in range(12):
    ____①____
max_cnt = cnt
for i in range(1, n - 11):
    cnt = cnt - flag[i - 1] + ____②____
    if cnt > max_cnt:
        max_cnt = cnt
        ____③____
''')
    para(doc, "答：①____________________  ②____________________  ③____________________")

    para(doc, "8. 用单调队列求列表 a 中每个长度为 k 的连续窗口最大值。请选择合适选项。", True)
    code(doc, '''
a = [4, 2, 5, 1, 6, 3]
k = 3
q = [0] * len(a)
head = tail = 0
ans = []
for i in range(len(a)):
    while head < tail and ______①______:
        tail -= 1
    q[tail] = i
    tail += 1
    if ______②______:
        head += 1
    if i >= k - 1:
        ans.append(______③______)
print(ans)
''')
    para(doc, "A. ① `a[q[tail-1]] <= a[i]`  ② `q[head] <= i-k`  ③ `a[q[head]]`")
    para(doc, "B. ① `a[q[tail-1]] >= a[i]`  ② `q[head] < i-k`  ③ `q[head]`")
    para(doc, "C. ① `a[q[head]] <= a[i]`  ② `q[tail-1] <= i-k`  ③ `a[q[tail]]`")
    para(doc, "D. ① `a[q[tail-1]] < a[i]`  ② `q[head] <= i-k+1`  ③ `a[i]`")
    para(doc, "答：______")

    doc.save(STUDENT_DOCX)


def build_answer():
    doc = new_doc()
    add_title(doc, "202606 九校期末错题专题提升卷（难题精选）答案", "教师用")
    section(doc, "答案")
    para(doc, "1. ① `maxv = ave`；② `t[5:7]`；③ B；④ F；⑤ C。")
    para(doc, "2. ① B；② C；③ F。")
    para(doc, "3. ① B；② D；③ F。")
    para(doc, "4. ① A；② C；③ F。")
    para(doc, "5. （1）2；① `total = 0`；② `lst[v] = tmp`；③ `ave > pt`。")
    para(doc, "6. （1）2；① `len(hum)-m+1`；② `bd <= T`；③ `state != pre_state`。")
    para(doc, "7. ① `cnt += flag[i]`；② `flag[i+11]`；③ `start_time = i`。")
    para(doc, "8. A。")
    doc.save(ANSWER_DOCX)


if __name__ == "__main__":
    OUT_DIR.mkdir(exist_ok=True)
    build_student()
    build_answer()
    print(STUDENT_DOCX)
    print(ANSWER_DOCX)
