from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT = Path("20出题/算法思维专题练习卷.docx")


def set_font(run, size=10.5, bold=False, name="宋体"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def para(doc, text="", first=False):
    text = text.replace("`", "")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if first:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_font(r)
    return p


def title(doc, text, sub=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, 16, True, "黑体")
    if sub:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(sub)
        set_font(r, 9)


def section(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, 12, True, "黑体")


def code(doc, text):
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        set_font(r, 9.2, False, "Consolas")


def add_img(doc, rel_path, width_cm=13.5):
    path = Path(rel_path)
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Cm(width_cm))


def new_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)
    return doc


def build():
    doc = new_doc()
    title(doc, "算法思维专题练习卷", "训练状态机、连续段、区间枚举、动态规划、调度模拟")
    para(doc, "班级：__________  姓名：__________  建议用时：60 分钟")

    section(doc, "一、状态机：限流与取消限流")
    para(doc, "1. 某车辆智能调度系统每分钟统计一次各通道的车辆通行数据，计算时刻在下一分钟一开始。当前车辆数 = 上一分钟一开始车辆数 + 上一分钟内的入场数 - 上一分钟内的离场数，并基于当前车辆数判定通道状态：", True)
    para(doc, "小于 20：状态码 0（通畅）；20～39：状态码 1（缓行）；大于等于 40：状态码 2（拥堵）。")
    para(doc, "调度规则：未限流时，若状态 1 连续持续 3 分钟，或状态变为 2，则立即发送一次“限流”指令；已限流时，若当前车辆数下降至 30 及以下，则立即发送一次“取消限流”指令；指令仅在条件满足时发送一次，不重复发送。")
    para(doc, "（1）若某通道在第 X+1 分钟一开始计算得到的当前车辆数为 15，接下来 6 分钟的入场数和离场数如下图，该通道首次发送“限流”指令是在第 X+______ 分钟一开始的时候。")
    add_img(doc, "attachments/202605县域教研高三月考_14_图1.png")
    para(doc, "（2）请在划线处填入合适代码。")
    code(doc, """
n = 4
k = [0] * n
code = [-1] * n
flag = [False] * n
s = [0] * n
while True:
    # q=[11,4,5,3,12,11,22,13] 表示各通道上一分钟入场数和离场数
    for i in range(n):
        _____①_____
        if s[i] < 20:
            code[i] = 0
        elif s[i] < 40:
            code[i] = 1
        else:
            code[i] = 2
        if code[i] == 1:
            _____②_____
        else:
            k[i] = 0
        if (code[i] == 2 or k[i] == 3) and not flag[i]:
            # 发送限流指令，代码略
            flag[i] = True
            k[i] = 0
        elif _____③_____:
            # 发送取消限流指令，代码略
            flag[i] = False
""")

    section(doc, "二、连续段容错：专注学习段")
    para(doc, "2. 每隔 5 分钟通过红外传感器检测一次状态：有人记为 1，无人记为 0。连续状态为 1 视为学习中；若学习中仅有 1 次检测为 0，视为“短暂休息”，不中断学习段；若第 2 次及以上为 0，则视为学习中断；当学习段中累计“有人”的次数 ≥ 4 时，判定为“专注学习段”。", True)
    para(doc, "（1）若某晚共检测 12 次，状态数据依次为：1，1，1，0，1，0，0，1，1，1，1，0，则“专注学习段”共有 ______ 个。")
    para(doc, "（2）请在划线处填入合适代码。")
    code(doc, """
# 读取状态数据存入列表 a，代码略
cnt = mx = c1 = 0
f = True
_____①_____
for i in range(n):
    if a[i] == 1:
        c1 += 1
    elif _____②_____:
        f = False
    else:
        if c1 >= 4:
            cnt += 1
            mx = max(mx, c1)
        f = True
        _____③_____
if c1 >= 4:
    cnt += 1
    mx = max(mx, c1)
print("专注学习段共有", cnt, "个", "最长持续", mx * 5, "分钟")
""")

    section(doc, "三、区间枚举：最长最佳时间段")
    para(doc, "3. 口袋公园通过摄像头与 AI 算法实时分析人流量。社区计划从历史人流量与温度数据中选择最佳活动时间段。若人流量处于 [L, R] 区间，且区间内温差不超过 5℃，即为最佳时间段。给定过往数据，找出最佳时段的最长长度。", True)
    para(doc, "（1）某天每小时人流量和平均温度数据如下图，要求人流量区间为 [15,35]，则活动可最长举办 ______ 小时。")
    add_img(doc, "attachments/202605强基联盟高三月考_14_图1.png")
    para(doc, "（2）请在划线处填入合适代码，并改正加框处错误条件。")
    code(doc, """
T = 5
L = 15
R = 35
mlen = 0
n = len(data)
for i in range(n):
    p = data[i][0]
    if L <= p <= R:
        _____①_____
        temps = []
        while j < n:
            flow, temp = data[j]
            if L <= flow <= R:   # 加框处代码有误
                break
            temps.append(temp)
            if max(temps) - min(temps) > T:
                break
            _____②_____
        mlen = max(mlen, j - i)
print(mlen)
""")
    para(doc, "加框处应改为：________________________________")

    section(doc, "四、动态规划：数字三角形最小路径和")
    para(doc, "4. 给定一个数字三角形，找出自顶向下的最小路径和。每一步只能移动到下一行中相邻的结点上。若某一步位于当前行下标 i 的结点，则下一步可移动到下一行下标 i 或 i+1 的结点。", True)
    add_img(doc, "attachments/202406宁波九校高二_13_图1.png", width_cm=12.5)
    para(doc, "（1）若将图中第四行第二列的数字 1 改为 7，则新的最小路径和为 ______。")
    para(doc, "（2）请在程序划线处填入合适代码。")
    code(doc, """
a = [[2], [3, 4], [6, 5, 7], [4, 1, 8, 3]]
n = len(a)
f = [[0] * n for i in range(n)]
f[0][0] = a[0][0]
for i in range(1, n):
    f[i][0] = f[i-1][0] + a[i][0]
    for j in range(1, i):
        if       ①        :
            f[i][j] = f[i-1][j-1] + a[i][j]
        else:
            f[i][j] = f[i-1][j] + a[i][j]
    f[i][i] =       ②
          ③
for i in range(1, n):
    if f[n-1][i] < ans:
        ans = f[n-1][i]
print("最小路径和为：", ans)
""")

    section(doc, "五、调度模拟：最少工人数")
    para(doc, "5. 某仓库需要安排工人搬运送达的货物。每辆货车记录有到达时间和货物箱数。所有货物必须在截止时间 T 内完成搬运。", True)
    para(doc, "规则：每个工人一次最多连续搬运 4 箱，每箱搬运耗时 3 单位时间；每次连续搬运结束后必须休息 1 单位时间；货车到达时，若存在已启用且空闲的工人，则优先安排其中最早空闲的工人，否则启用新工人。已启用工人总数不能超过上限 n。")
    para(doc, "（1）若货车记录列表 data=[[1,5],[9,7],[17,6]]，T=30，则需要的工人数量为 ______。")
    para(doc, "（2）定义 wsort(data) 函数，根据到达时间升序排序。划线处应填入的代码为 ______（多选，填字母）。")
    code(doc, """
def wsort(data):
    for i in range(len(data)-1):
        for j in range(_____________):
            if data[j][0] > data[j+1][0]:
                data[j], data[j+1] = data[j+1], data[j]
    return data
""")
    para(doc, "A. len(data)-i-1    B. len(data)-i    C. len(data)-2,i,-1    D. len(data)-2,i-1,-1")
    para(doc, "（3）请在划线处填入合适代码。")
    code(doc, """
def get_min(lst):
    if len(lst) == 0:
        return -1
    min_idx = 0
    for i in range(1, len(lst)):
        if lst[i] < lst[min_idx]:
            min_idx = i
    return min_idx

def calc(data, T, n):
    max_r = 4; dt = 3; count = 0; ft = 1; worker = []
    for task in data:
        st, box = task[0], task[1]
        remain = T - st - ft
        if remain < dt:
            return -1
        while box > 0:
            use_new = True
            idx = get_min(worker)
            if ____①____:
                worker = worker[:idx] + worker[idx+1:]
                use_new = False
            if use_new:
                count += 1
                if count > n:
                    return -1
            once = min(remain // dt, box, max_r)
            new_time = ____②____
            worker.append(new_time)
            box -= once
    return count

data = wsort(data)
ans = calc(data, T, n)
if ____③____:
    print("任务无法完成！")
else:
    print("最少需要", ans, "名工人！")
""")

    doc.add_page_break()
    title(doc, "参考答案")
    para(doc, "1. （1）3；（2）① s[i] = s[i] + q[2*i] - q[2*i+1]；② k[i] += 1；③ flag[i] and s[i] <= 30。")
    para(doc, "2. （1）2；（2）① n = len(a)；② c1 > 0 and f；③ c1 = 0。")
    para(doc, "3. （1）3；（2）① j = i；② j += 1；加框处改为 flow < L or flow > R。")
    para(doc, "4. （1）15；（2）① f[i-1][j-1] < f[i-1][j]；② f[i-1][i-1] + a[i][i]；③ ans = f[n-1][0]。")
    para(doc, "5. （1）4；（2）A、D；（3）① len(worker) > 0 and worker[idx] <= st；② st + once * dt + ft；③ ans == -1。")

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
